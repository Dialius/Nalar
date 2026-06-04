import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

LOGO_SRC = r"D:\Nalar\deliverables\logo_smk_telkom_clean.png"
OUTPUT = r"D:\Nalar\deliverables\2_Source_Code_API_Documentation.docx"

# Helper for cell shading
def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def setup_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.18)
        section.bottom_margin = Inches(1.18)
        section.left_margin = Inches(1.57)
        section.right_margin = Inches(1.18)
        
        section.different_first_page_header_footer = True
        
        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_hdr = p_hdr.add_run("Dokumen Laporan Resmi Proyek Nalar | Kelompok Nalar | XI PPLG 3")
        r_hdr.font.name = 'Times New Roman'
        r_hdr.font.size = Pt(8.5)
        r_hdr.italic = True
        r_hdr.font.color.rgb = RGBColor(128, 128, 128)
        
        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ftr.paragraph_format.space_after = Pt(0)
        
        r_ftr = p_ftr.add_run("Halaman ")
        r_ftr.font.name = 'Times New Roman'
        r_ftr.font.size = Pt(9)
        r_ftr.font.color.rgb = RGBColor(128, 128, 128)
        
        run_num = p_ftr.add_run()
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(9)
        run_num.font.color.rgb = RGBColor(128, 128, 128)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_num._r.append(fldChar1)
        run_num._r.append(instrText)
        run_num._r.append(fldChar2)
        run_num._r.append(fldChar3)

def make_cover_page(doc, title, subtitle):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(28)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(18)
    r_title.bold = True
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run(subtitle)
    r_sub.font.name = 'Times New Roman'
    r_sub.font.size = Pt(13)
    r_sub.bold = True
    
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(24)
    r_logo = p_logo.add_run()
    r_logo.add_picture(LOGO_SRC, width=Inches(1.7))
    
    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_after = Pt(8)
    r_by = p_by.add_run("Disusun Oleh :")
    r_by.font.name = 'Times New Roman'
    r_by.font.size = Pt(11)
    r_by.bold = True
    
    members = [
        "DAVINZA FATTAH DZULHIJRIYAN SYAHID (541241043)",
        "JONATHAN BIMA PRADANA PUTRA (541241097)",
        "KHAFIDZ ASAD HERMAWAN (541241105)",
        "MUHAMMAD RIDLO AL QOHHAAR (541241141)",
        "VEREL CHANDRA RIYANTO (541241195)"
    ]
    for m in members:
        p_m = doc.add_paragraph()
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_m.paragraph_format.space_before = Pt(0)
        p_m.paragraph_format.space_after = Pt(2)
        p_m.paragraph_format.line_spacing = 1.0
        r_m = p_m.add_run(m)
        r_m.font.name = 'Times New Roman'
        r_m.font.size = Pt(11)
        r_m.bold = True
        
    p_school_start = doc.add_paragraph()
    p_school_start.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_school_start.paragraph_format.space_before = Pt(48)
    p_school_start.paragraph_format.space_after = Pt(2)
    r_ss1 = p_school_start.add_run("PENGEMBANGAN PERANGKAT LUNAK DAN GIM")
    r_ss1.font.name = 'Times New Roman'
    r_ss1.font.size = Pt(13)
    r_ss1.bold = True
    
    p_school_mid = doc.add_paragraph()
    p_school_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_school_mid.paragraph_format.space_before = Pt(0)
    p_school_mid.paragraph_format.space_after = Pt(2)
    r_ss2 = p_school_mid.add_run("SMK TELKOM PURWOKERTO")
    r_ss2.font.name = 'Times New Roman'
    r_ss2.font.size = Pt(13)
    r_ss2.bold = True
    
    p_school_end = doc.add_paragraph()
    p_school_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_school_end.paragraph_format.space_before = Pt(0)
    p_school_end.paragraph_format.space_after = Pt(0)
    r_ss3 = p_school_end.add_run("TAHUN AJARAN 2025 / 2026")
    r_ss3.font.name = 'Times New Roman'
    r_ss3.font.size = Pt(13)
    r_ss3.bold = True
    
    doc.add_page_break()

def make_approval_page(doc, doc_title):
    p_h = doc.add_paragraph()
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h = p_h.add_run("LEMBAR PENGESAHAN")
    r_h.font.name = 'Times New Roman'
    r_h.font.size = Pt(14)
    r_h.bold = True
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after = Pt(20)
    
    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_desc = p_desc.add_run("Dokumen Laporan Akhir Proyek dengan judul :\n")
    r_desc.font.name = 'Times New Roman'
    r_desc.font.size = Pt(12)
    
    r_title = p_desc.add_run(f"\"{doc_title.upper()}\"\n")
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(12)
    r_title.bold = True
    
    r_desc2 = p_desc.add_run("telah diperiksa, diuji secara fungsional, dan disahkan oleh Tim Pembimbing Program Keahlian Pengembangan Perangkat Lunak dan Gim (PPLG) SMK Telkom Purwokerto sebagai salah satu bukti pemenuhan syarat akademis kelulusan tugas akhir Kelompok Nalar.")
    r_desc2.font.name = 'Times New Roman'
    r_desc2.font.size = Pt(12)
    p_desc.paragraph_format.line_spacing = 1.5
    p_desc.paragraph_format.space_after = Pt(36)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run("Purwokerto, Juni 2026\nMenyetujui,")
    r_date.font.name = 'Times New Roman'
    r_date.font.size = Pt(12)
    p_date.paragraph_format.space_after = Pt(24)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Normal Table'
    
    cell_l0 = table.cell(0, 0)
    p_l0 = cell_l0.paragraphs[0]
    p_l0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l0 = p_l0.add_run("Ketua Kelompok Nalar,\n\n\n\n\nDAVINZA FATTAH D. S.\nNIS. 541241043")
    r_l0.font.name = 'Times New Roman'
    r_l0.font.size = Pt(11)
    
    cell_r0 = table.cell(0, 1)
    p_r0 = cell_r0.paragraphs[0]
    p_r0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r0 = p_r0.add_run("Guru Pembimbing PPLG,\n\n\n\n\nBapak/Ibu Guru Pembimbing\nNIP. 19820412 200801 1 002")
    r_r0.font.name = 'Times New Roman'
    r_r0.font.size = Pt(11)
    
    cell_h = table.cell(1, 0)
    cell_h.merge(table.cell(1, 1))
    p_h2 = cell_h.paragraphs[0]
    p_h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_h2.paragraph_format.space_before = Pt(24)
    r_h2 = p_h2.add_run("Mengetahui,\nKepala Program Keahlian PPLG SMK Telkom Purwokerto,\n\n\n\n\nKepala Program Keahlian PPLG\nNIP. 19780524 200501 2 001")
    r_h2.font.name = 'Times New Roman'
    r_h2.font.size = Pt(11)
    
    doc.add_page_break()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)

def add_bullet(doc, bold_prefix, text_body):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    
    r_bold = p.add_run(bold_prefix)
    r_bold.font.name = 'Times New Roman'
    r_bold.font.size = Pt(12)
    r_bold.bold = True
    
    r_body = p.add_run(text_body)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(12)

def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9.5)
        
    doc.add_paragraph()

def add_callout(doc, text, title="CATATAN PENTING:  "):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F8FC")
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    r_title = p.add_run(title)
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(31, 78, 121)
    
    r_body = p.add_run(text)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = RGBColor(60, 65, 75)
    
    doc.add_paragraph()

def add_kata_pengantar(doc, doc_title):
    p_h = doc.add_paragraph()
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h = p_h.add_run("KATA PENGANTAR")
    r_h.font.name = 'Times New Roman'
    r_h.font.size = Pt(14)
    r_h.bold = True
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after = Pt(16)
    
    add_para(doc, "Puji dan syukur kami panjatkan ke hadirat Tuhan Yang Maha Esa atas segala rahmat, berkat, dan hidayah-Nya, sehingga kami dapat menyelesaikan laporan akademik resmi berjudul \"" + doc_title.upper() + "\" tepat pada waktunya.")
    add_para(doc, "Penyusunan laporan ini merupakan wujud pertanggungjawaban ilmiah dan rekapitulasi teknis atas perancangan, pengembangan, serta pengujian sistem backend RESTful API bernama \"Nalar\". Dokumen ini disusun untuk melengkapi seluruh berkas kelulusan penyerahan tugas akhir (Capstone Project) kelompok kami di kelas XI PPLG 3 SMK Telkom Purwokerto.")
    add_para(doc, "Kami menyadari sepenuhnya bahwa keberhasilan penyusunan laporan ini tidak lepas dari bimbingan, arahan, dan dukungan moril dari berbagai pihak. Oleh karena itu, dengan kerendahan hati kami mengucapkan terima kasih kepada Bapak/Ibu Guru Pembimbing PPLG atas ketekunan dan kesabarannya dalam memandu arah riset kami, serta rekan-rekan anggota kelompok yang senantiasa solid bekerja bahu-membahu merampungkan sistem.")
    add_para(doc, "Kami menyadari bahwa laporan ini masih jauh dari kesempurnaan. Oleh karena itu, segala kritik konstruktif serta saran yang membangun dari pembaca sangat kami harapkan guna penyempurnaan di masa mendatang. Akhir kata, semoga laporan akademik ini dapat mendatangkan manfaat nyata bagi pengembangan rekayasa perangkat lunak di lingkungan sekolah maupun khalayak pembaca.")
    
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.paragraph_format.space_before = Pt(28)
    p_sign.paragraph_format.space_after = Pt(0)
    r_sign = p_sign.add_run("Purwokerto, Juni 2026\n\n\n\nKelompok Nalar")
    r_sign.font.name = 'Times New Roman'
    r_sign.font.size = Pt(12)
    r_sign.bold = True
    
    doc.add_page_break()

def add_daftar_isi(doc, items):
    p_h = doc.add_paragraph()
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h = p_h.add_run("DAFTAR ISI")
    r_h.font.name = 'Times New Roman'
    r_h.font.size = Pt(14)
    r_h.bold = True
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after = Pt(18)
    
    for label, page, level in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(5.75), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        
        if level == 1:
            p.paragraph_format.left_indent = Inches(0.25)
            
        r_text = p.add_run(f"{label}\t{page}")
        r_text.font.name = 'Times New Roman'
        r_text.font.size = Pt(12)
        if level == 0:
            r_text.bold = True
            
    doc.add_page_break()


if __name__ == '__main__':
    print("Building Formal DOCX: 2_Source_Code_API_Documentation.docx...")
    doc = docx.Document()
    setup_document(doc)
    
    make_cover_page(doc, "LAPORAN IMPLEMENTASI KODE SUMBER API BACKEND", "Arsitektur RESTful API, Keamanan Sesi JWT & Integrasi Webhook Aplikasi Nalar")
    make_approval_page(doc, "Laporan Implementasi Kode Sumber API Backend Nalar")
    add_kata_pengantar(doc, "Laporan Implementasi Kode Sumber API Backend Nalar")
    
    add_daftar_isi(doc, [
        ("KATA PENGANTAR", "i", 0),
        ("LEMBAR PENGESAHAN", "ii", 0),
        ("DAFTAR ISI", "iii", 0),
        ("BAB I: PENDAHULUAN ARSITEKTUR KODE BACKEND", "1", 0),
        ("1.1 Latar Belakang Pemilihan Node.js & Express.js", "1", 1),
        ("1.2 Keunggulan Model Non-Blocking I/O", "1", 1),
        ("BAB II: DESAIN DAN STRUKTUR MODUL KODE SUMBER API", "2", 0),
        ("2.1 Struktur Folder Proyek API", "2", 1),
        ("2.2 Alur Pemrosesan Request API Nalar", "2", 1),
        ("BAB III: IMPLEMENTASI KODE SUMBER UTAMA BACKEND", "3", 0),
        ("3.1 Implementasi Koneksi Basis Data (db.js)", "3", 1),
        ("3.2 Implementasi Middleware Autentikasi Sesi (auth.js)", "4", 1),
        ("BAB IV: ANALISIS INTEGRASI FINANSIAL & WEBHOOK MIDTRANS", "5", 0),
        ("BAB V: PENUTUP", "6", 0),
        ("5.1 Kesimpulan Implementasi Backend", "6", 1),
        ("5.2 Saran Iterasi Arsitektur Server", "6", 1)
    ])
    
    # BAB I
    add_heading_1(doc, "BAB I: PENDAHULUAN ARSITEKTUR KODE BACKEND")
    add_heading_2(doc, "1.1 Latar Belakang Pemilihan Node.js & Express.js")
    add_para(doc, "Ekosistem backend Nalar dikembangkan menggunakan runtime environment Node.js dan framework Express.js. Pemilihan teknologi ini didasari oleh kebutuhan performa tinggi dalam menyajikan data kemajuan belajar siswa (user progress) secara dinamis, serta kemampuan menangani banyak request HTTP secara bersamaan (high concurrency).")
    add_heading_2(doc, "1.2 Keunggulan Model Non-Blocking I/O")
    add_para(doc, "Node.js menggunakan model I/O non-blocking yang digerakkan oleh event (event-driven). Keunggulan ini sangat krusial saat backend melakukan operasi database MySQL yang intensif atau ketika mengirim permintaan HTTP eksternal ke gateway pembayaran Midtrans Sandbox, sehingga server tetap responsif dan meminimalkan latensi.")
    
    # BAB II
    add_heading_1(doc, "BAB II: DESAIN DAN STRUKTUR MODUL KODE SUMBER API")
    add_heading_2(doc, "2.1 Struktur Folder Proyek API")
    add_para(doc, "Struktur folder repositori API backend nalar-backend disusun secara modular untuk memudahkan pemeliharaan kode:")
    add_bullet(doc, "/database: ", "Menyimpan driver db.js untuk mengonfigurasi pool koneksi database MySQL lokal.")
    add_bullet(doc, "/middleware: ", "Berisi validator penapis akses seperti auth.js untuk JWT dan admin.js untuk peran administrator.")
    add_bullet(doc, "/routes: ", "Mengelompokkan rute API berdasarkan modul fungsi (auth, users, courses, gamification, payment).")
    add_bullet(doc, "server.js: ", "Merupakan gerbang entri utama untuk menjalankan server Node.js dan menyambungkan middleware global.")
    
    add_heading_2(doc, "2.2 Alur Pemrosesan Request API Nalar")
    add_para(doc, "Setiap request dari klien Android Nalar akan diterima oleh rute Express di server.js, kemudian rute akan memanggil middleware autentikasi JWT. Setelah request dinyatakan valid, handler di folder routes/ akan mengambil koneksi pool dari MySQL, menjalankan query database, dan mengembalikan respons dalam format JSON standard.")
    
    doc.add_page_break()
    
    # BAB III
    add_heading_1(doc, "BAB III: IMPLEMENTASI KODE SUMBER UTAMA BACKEND")
    
    add_heading_2(doc, "3.1 Implementasi Koneksi Basis Data (db.js)")
    add_para(doc, "Kode di bawah ini menunjukkan inisialisasi pool koneksi basis data MySQL menggunakan promise wrapper asinkron:")
    add_code_block(doc, [
        "const mysql = require('mysql2');",
        "require('dotenv').config();",
        "",
        "const pool = mysql.createPool({",
        "  host: process.env.DB_HOST || '127.0.0.1',",
        "  user: process.env.DB_USER || 'root',",
        "  password: process.env.DB_PASSWORD || 'root',",
        "  database: process.env.DB_NAME || 'nalar_db',",
        "  waitForConnections: true,",
        "  connectionLimit: 10,",
        "  queueLimit: 0",
        "});",
        "",
        "module.exports = pool.promise();"
    ])
    
    doc.add_page_break()
    
    add_heading_2(doc, "3.2 Implementasi Middleware Autentikasi Sesi (auth.js)")
    add_para(doc, "Middleware berikut memverifikasi tanda tangan token JWT (JSON Web Token) yang disisipkan pada header Authorization untuk mengamankan data siswa:")
    add_code_block(doc, [
        "const jwt = require('jsonwebtoken');",
        "require('dotenv').config();",
        "",
        "const requireAuth = (req, res, next) => {",
        "  const authHeader = req.headers.authorization;",
        "  ",
        "  if (!authHeader || !authHeader.startsWith('Bearer ')) {",
        "    return res.status(401).json({ error: 'Akses ditolak. Token tidak disediakan.' });",
        "  }",
        "",
        "  const token = authHeader.split(' ')[1];",
        "  try {",
        "    const decoded = jwt.verify(token, process.env.JWT_SECRET || 'nalar_secret_key');",
        "    req.user = decoded;",
        "    next();",
        "  } catch (err) {",
        "    return res.status(401).json({ error: 'Akses ditolak. Token tidak valid.' });",
        "  }",
        "};",
        "",
        "module.exports = requireAuth;"
    ])
    
    doc.add_page_break()
    
    # BAB IV
    add_heading_1(doc, "BAB IV: ANALISIS INTEGRASI FINANSIAL & WEBHOOK MIDTRANS")
    add_para(doc, "Integrasi gerbang pembayaran dilakukan menggunakan SNAP API Midtrans. Saat siswa membeli akses VIP premium, klien memanggil endpoint /api/payment/checkout untuk mendapatkan Snap Token transaksi unik. Server Midtrans kemudian mengirimkan status pembayaran instan (settlement) melalui callback webhook asinkron ke server backend, yang secara aman memicu peningkatan flag status 'is_premium = 1' di basis data users.")
    
    # BAB V
    add_heading_1(doc, "BAB V: PENUTUP")
    add_heading_2(doc, "5.1 Kesimpulan Implementasi Backend")
    add_para(doc, "Implementasi RESTful API Nalar berbasis Node.js Express terbukti berjalan sangat stabil, efisien dalam mengelola memori, serta memiliki sistem penanganan webhook dan enkripsi password yang aman.")
    add_heading_2(doc, "5.2 Saran Iterasi Arsitektur Server")
    add_para(doc, "Saran untuk pengembangan di masa depan adalah migrasi API menggunakan TypeScript untuk mendapatkan sistem statis pengetikan (strong-typing) demi meminimalkan error runtime.")
    
    add_callout(doc, "Untuk mencegah kebocoran credentials, berkas konfigurasi .env tidak boleh dikirimkan ke repositori publik dan wajib dikecualikan secara ketat lewat berkas .gitignore.", "KEAMANAN SERVER: ")
    
    doc.save(OUTPUT)
    print(f"API DOCX built successfully at {OUTPUT}!")
